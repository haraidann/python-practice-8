from __future__ import annotations
import os
from typing import Dict, Any, Optional
from jinja2 import Environment, FileSystemLoader, select_autoescape, Template
from weasyprint import HTML
from docx import Document 
from docx.shared import Pt
import qrcode
from io import BytesIO
from PIL import Image
import datetime
from quest_master.core.database import Database


class TemplateEngine:
    def __init__(self, templates_dir: Optional[str] = None):
        if templates_dir:
            base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            full_path = os.path.join(base_dir, templates_dir)
            loader = FileSystemLoader(full_path)
            self.env = Environment(
                loader=loader,
                autoescape=select_autoescape(["html", "xml"])
            )
        else:
            self.env = Environment(autoescape=select_autoescape(["html", "xml"]))

    def __init__(self, templates_dir: Optional[str] = None):
        if templates_dir:
            loader = FileSystemLoader(templates_dir)
            self.env = Environment(
                loader=loader,
                autoescape=select_autoescape(["html", "xml"])
            )
        else:
            self.env = Environment(autoescape=select_autoescape(["html", "xml"]))

    def render_from_string(self, template_str: str, context: Dict[str, Any]) -> str:
        tpl: Template = self.env.from_string(template_str)
        return tpl.render(**context)

    def render_from_file(self, template_name: str, context: Dict[str, Any]) -> str:
        tpl = self.env.get_template(template_name)
        return tpl.render(**context)

    @staticmethod
    def generate_qr(data: str, output_path: Optional[str] = None, box_size: int = 10) -> bytes:

        qr = qrcode.QRCode(version=1, box_size=box_size, border=2)
        qr.add_data(data)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")

        bio = BytesIO()
        img.save(bio, format="PNG")
        png_bytes = bio.getvalue()
        if output_path:
            with open(output_path, "wb") as f:
                f.write(png_bytes)
        return png_bytes

    @staticmethod
    def html_to_pdf(html_str: str, output_path: str, base_url: Optional[str] = None) -> None:

        HTML(string=html_str, base_url=base_url).write_pdf(output_path)


    @staticmethod
    def render_to_docx_from_text(text: str, output_path: str, title_style: bool = True) -> None:
        doc = Document()
        lines = text.splitlines()
        if title_style and lines:
            for i, line in enumerate(lines):
                if line.strip():
                    doc.add_heading(line.strip(), level=1)
                    remaining = "\n".join(lines[i + 1 :]).strip()
                    if remaining:
                        for para in remaining.split("\n\n"):
                            p = doc.add_paragraph(para.strip())
                            for run in p.runs:
                                run.font.size = Pt(11)
                    break
        else:
            for para in "\n\n".join(lines).split("\n\n"):
                p = doc.add_paragraph(para.strip())
                for run in p.runs:
                    run.font.size = Pt(11)

        doc.save(output_path)

    def render_context_to_pdf(self, template_str: str, context: Dict[str, Any], output_path: str,
                              embed_qr: Optional[str] = None, base_url: Optional[str] = None) -> None:

        if embed_qr:
            png_bytes = self.generate_qr(embed_qr)
            import base64
            b64 = base64.b64encode(png_bytes).decode("ascii")
            data_uri = f"data:image/png;base64,{b64}"
            context = dict(context)
            context["qr_img_data"] = data_uri

        html = self.render_from_string(template_str, context)
        self.html_to_pdf(html, output_path, base_url=base_url)

    def render_context_to_docx(self, template_str: str, context: Dict[str, Any], output_path: str) -> None:
        text = self.render_from_string(template_str, context)
        self.render_to_docx_from_text(text, output_path)


class BatchExporter:
    @staticmethod
    def generate_100_quests(db: Database, te: TemplateEngine, output_dir: str = "batch/"):
        import os
        os.makedirs(output_dir, exist_ok=True)
        for i in range(100):
            title = f"Batch Quest {i}"
            quest_id = db.create_quest(title, "Легкий", 10, "Описание " * 50, "2025-12-31")
            ctx = {"quest": db.get_quest(quest_id), "now": datetime.datetime.now().isoformat()}
            html = te.render_from_file("royal_decree.html", ctx)
            te.html_to_pdf(html, f"{output_dir}/quest_{i}.pdf")

